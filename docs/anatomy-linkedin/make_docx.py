from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- CONTENT BLOCKS (each one is copy-paste-ready, character-counted) ----------

HOOK_140 = (
    "We shipped an AI-governance page that recomputes a SHA-256 over the "
    "published thesis bundle on every load. Drifts → red banner."
)  # target ≤ 140 (desktop preview)

HOOK_210 = (
    "Most \"AI governance\" pages are decoration. Ours recomputes a SHA-256 "
    "over the entire upstream-published anatomy bundle on every page load. "
    "One byte drifts, the banner flips red. No trust-us path."
)  # target ≤ 210 (mobile fold)

POST_FULL = """Most "AI governance" pages in production are decoration: a Notion link, a PDF on a shared drive, a screenshot of a SOC 2 badge. We just shipped the opposite.

A11oy → /anatomy is a read-only viewer that vendors the entire upstream Ouroboros thesis anatomy bundle (CC-BY-4.0, github.com/szl-holdings/ouroboros-thesis) and proves byte-level parity with the published thesis on every page load. If one byte drifts, the banner flips from green to red and the page refuses to claim parity. There is no "trust us" code path.

WHAT'S IN THE BUNDLE

• 8 canonical figures (brain, wires, full-body, heart, blood/immune, skeleton, nervous, body-graph) as both vector PDF and 300dpi PNG — the operational schematic of a governed agent: AMARU cortex with a 9-axis conjunctive doctrine gate, YUYAY v3 13-axis AND gate (no averaging, sacred axes ≥ 0.95), YAWAR append-only ledger, HUKLLA T01–T10 tripwires, OTel/VSP nervous system, CHAKANA 21-edge M=0 lattice.

• 7-chakra activation spine (root → crown) with per-chakra leader.md + result.md prose pulled verbatim from upstream — the actual kernel notes (tinygrad → KALLPA dispatch, DSPy → YACHAY retrieval, vLLM → RIMAY sampling), not marketing copy.

• Real DOIs on every section: concept DOI 10.5281/zenodo.19944926 (always resolves to latest), v13 release 10.5281/zenodo.20195368.

THE INTEGRITY MECHANISM (the part CTOs care about)

In the browser, on every load:

1. Fetch VENDOR.json::expected_files (canonical 14-file manifest).
2. For each filename in sorted order, fetch the bytes and concat sha256( filename || NUL || bytes || NUL ).
3. Compare the digest to VENDOR.json::upstream_sha. Match → green. Mismatch → red DRIFT banner with both hashes side-by-side. No SubtleCrypto → sentinel that never matches a real sha256. Fails closed, never silently green.

Same construction as Sigstore cosign verify or Subresource Integrity, applied to a documentation bundle.

WHY THIS MATTERS OPERATIONALLY

• Auditors re-run the hash and get bit-for-bit reproducibility against the publication.
• CI can fail the build the moment someone "fixes a typo" in a vendored binary.
• Citations are stable — every chakra section deep-links to a resolvable Zenodo DOI, not a Notion page that 404s in six months.
• If you're claiming "auditable AI," these eight figures are the load-bearing primitives. Run a hash against your own published claims and see what happens.

STACK

React + Vite + TypeScript. PDF via <object> with PNG fallback (CSP-aware). crypto.subtle.digest("SHA-256", …). No third-party hash libs — the proof is ~30 lines and inspectable in DevTools.

Upstream thesis: github.com/szl-holdings/ouroboros-thesis"""

POST_SHORT = """Most "AI governance" pages are decoration. Ours isn't.

A11oy → /anatomy vendors the entire upstream Ouroboros thesis anatomy bundle (CC-BY-4.0) and recomputes a SHA-256 over the published binaries in the browser on every page load. One byte drifts, the banner flips red. No "trust us" path.

In the bundle:
• 8 canonical figures (brain, wires, full-body, heart, blood/immune, skeleton, nervous, body-graph) — the operational schematic of a governed agent: AMARU cortex (9-axis doctrine gate), YUYAY v3 (13-axis AND gate, no averaging), YAWAR append-only ledger, HUKLLA tripwires, OTel/VSP nervous system.
• 7-chakra activation spine with per-kernel prose pulled verbatim from upstream.
• Real DOIs (concept 10.5281/zenodo.19944926, v13 10.5281/zenodo.20195368).

The integrity proof is ~30 lines of crypto.subtle and inspectable in DevTools. Same construction as Sigstore cosign verify, applied to a documentation bundle. Auditors re-run the hash; CI fails the build the moment a vendored byte changes.

If your governance claim can't survive a hash check, it's a brochure.

Upstream: github.com/szl-holdings/ouroboros-thesis"""

HASHTAGS = "#AIGovernance #AuditableAI #DecisionIntelligence #ReproducibleResearch #PlatformEngineering #SoftwareReceipts #SRE #OuroborosThesis #SZLHoldings"

WARHACKER_BODY = """TO: Defense Unicorns — Warhacker (defenseunicorns.com/warhacker)
RE: Fixing the drones — accountable autonomy on small UAS
FROM: SZL Holdings · CC-BY-4.0 · DOI 10.5281/zenodo.19944926

PITCH

Small-UAS autonomy isn't an airframe problem — it's an accountable-software problem. Closed binaries, no provenance, models that can't be re-run bit-for-bit after an engagement, "AI" decisions with no receipt. We built the substrate that fixes that — CC-BY-4.0, permanent DOI, drops into UDS / Big Bang, running live now. A JAG-defensible flight-data recorder for every kernel decision a drone makes.

WHAT WE JUST SHIPPED

/anatomy now embeds a live ops panel polling our amaru sidecar over a read-only proxy: chakras registered, receipts counter, scheduler ticks, bus publishes/failures, HUKLLA T01–T10 with pass/warn/trip dots, R0513 invariants + live kernel/brain hashes, chakana 21-edge lattice with the ouroboros edge in gold. Nothing is mocked — counters increment, tripwires flip, the cycle is visibly the only cycle. Same proxy + allowlist exposes drone-side state to a UDS dashboard with zero new auth. We also pulled the decorative AI hero images. Page shows only what we can prove.

HOW IT FIXES THE DRONES

1. Supply-chain integrity. SHA-256 recomputed over every vendored binary on load; one drifted byte and the banner flips red. Sigstore-class proof on the doctrine bundle. Ports as-is into Big Bang / UDS.

2. Doctrine gate (YUYAY v3). 13-axis conjunctive AND — sacred axes (moralGrounding, measurabilityHonesty, ontologicalGrounding) ≥0.95, rest ≥0.90. No averaging. Technical implementation of DoD 3000.09 human-judgment.

3. YAWAR receipts. Every evaluation appends a hash-chained receipt (seq, prev_hash, self_hash, params, result). Flight-data recorder for autonomy — replay the exact (input, model, doctrine) after any engagement.

4. HUKLLA tripwires. T01–T10 cover registration, chain integrity, scheduler progress, wiring acyclicity, bus health, doctrine load, envelope size, proof freshness. Trip → halt. Stops the fleet on a quiet model swap or stale pin.

5. R0513 OVERWATCH. Six invariants over the live receipt chain (KL drift, joint margin, mid-exec regate, M=0 rigidity, hash integrity, reserved). Read-only — halt stays with HUKLLA, so OVERWATCH is never an SPoF.

6. Byte-identical replay. Each kernel is ≤10 lines and ships a 5× byte-identical replay SHA-256. After-action review is `python3 test_replay.py` and a hash compare.

7. Minimization. Kernels are tiny (3.45% absorbed source, 50× reduction). Reviewable code, small SBOM, real audit pre-flight.

WHAT TO DO WITH IT

Vendor the bundle into your UDS / Big Bang package (hash self-check ports as-is). Put YUYAY v3 in front of any model emitting an action recommendation. Pipe autonomy decisions through YAWAR. Wire HUKLLA as CI gates AND runtime kill-switches — same definition both places. Adopt byte-identical replay for every kernel in an effects loop. If it can't replay, it can't fly.

Nothing here is a slide."""

FIRST_COMMENT = """Direct links if anyone wants to dig in:

• Upstream thesis (source of truth): github.com/szl-holdings/ouroboros-thesis
• Concept DOI (always latest): doi.org/10.5281/zenodo.19944926
• v13 release DOI: doi.org/10.5281/zenodo.20195368
• Anatomy bundle index (per-figure metadata + sha256 list): docs/anatomy/INDEX.md and docs/anatomy/figures.sha256 in the repo above

Happy to walk through the in-browser integrity check construction if anyone wants the 30-line walkthrough."""

# ---------- DOC ----------

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

def para(text, bold=False, italic=False, mono=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if mono:
        r.font.name = 'Consolas'
    return p

def block(text, label, limit_label):
    n = len(text)
    heading(f"{label} — {n} chars {limit_label}", level=2)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    doc.add_paragraph()  # spacer

title = doc.add_heading('LinkedIn post — A11oy /anatomy (CTO-facing, copy-paste pack)', level=0)
for r in title.runs:
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

para(
    'Each block below is plain text, ready to copy directly into LinkedIn. '
    'LinkedIn rules of thumb: ~140 chars visible in the desktop feed preview, '
    '~210 chars on mobile before the "see more" cut, 3,000 chars hard limit '
    'on a post, and the first comment is the place to drop links so the '
    'post body stays clean. Character counts are shown in each section heading.',
    italic=True,
)

heading('1. Above-the-fold hook (desktop preview, ≤140 chars)', 1)
para(
    'Use this as the opening line of the post so what appears before "see more" '
    'in the desktop feed is the kill-shot. Then paste the FULL POST body below it.',
    italic=True,
)
block(HOOK_140, 'HOOK — desktop preview', '(target ≤ 140)')

heading('2. Above-the-fold hook (mobile, ≤210 chars)', 1)
para(
    'Slightly longer alternative if you care more about the mobile feed than '
    'the desktop feed.',
    italic=True,
)
block(HOOK_210, 'HOOK — mobile fold', '(target ≤ 210)')

heading('3. FULL POST BODY (≤3,000 chars — LinkedIn hard limit)', 1)
para(
    'This is the main, technical, CTO-facing post. Plain text, no markdown — '
    'LinkedIn does not render Markdown. Bullet glyphs (•) and arrows (→) '
    'render fine. Triple-click the block below to select, then copy.',
    italic=True,
)
block(POST_FULL, 'FULL POST', '(LinkedIn limit 3,000)')

heading('4. SHORT POST (≤1,300 chars — higher reach variant)', 1)
para(
    'Use this if you want maximum dwell time / completion rate rather than '
    'depth. LinkedIn\'s algorithm tends to reward posts read all the way through, '
    'and shorter posts complete more often.',
    italic=True,
)
block(POST_SHORT, 'SHORT POST', '(target ≤ 1,300)')

doc.add_page_break()
heading('5. WARHACKER appendix — how this fixes the drones', 1)
para(
    'Long-form, CTO/PM-facing explanation of how the anatomy bundle maps onto '
    'Defense Unicorns\' Warhacker push for accountable software on small UAS. '
    'Plain text, copy-paste-ready into a LinkedIn article, a UDS / Big Bang '
    'design doc, or an email to a program manager. References real Warhacker '
    'context (defenseunicorns.com/warhacker) and concrete bundle primitives.',
    italic=True,
)
block(WARHACKER_BODY, 'WARHACKER APPENDIX', '(article-length, no LinkedIn 3,000 limit)')
doc.add_page_break()

heading('6. Hashtag block (paste at end of post or as first comment)', 1)
block(HASHTAGS, 'HASHTAGS', '')

heading('7. First comment (link drop — keeps body clean)', 1)
para(
    'LinkedIn deprioritizes posts with external links in the body. Standard '
    'practice: omit the GitHub/Zenodo URLs from the body and drop them as the '
    'first comment within ~60 seconds of posting.',
    italic=True,
)
block(FIRST_COMMENT, 'FIRST COMMENT', '')

doc.add_page_break()

heading('Suggested image stack (carousel order)', 1)
para('1. screenshots/a11oy_anatomy_top.jpg — green "Bundle integrity verified" banner. This is the kill-shot.')
para('2. agent-anatomy/anatomy_brain.png — AMARU cortex / 9-axis doctrine gate.')
para('3. agent-anatomy/anatomy_heart.png — YUYAY v3 13-axis conjunctive AND gate.')
para('4. agent-anatomy/anatomy_skeleton.png — 12-repo service topology.')
para('5. agent-anatomy/anatomy_nervous.png — OTel/VSP span propagation.')
para('6. agent-anatomy/anatomy_body_graph.png — master overlay (the "this is the system" hero).')

heading('Posting tips', 1)
para('• Post Tue/Wed/Thu between 8–10am in your audience\'s timezone for engineering audiences.')
para('• First 60 seconds matter: drop the first comment with links, like your own post once, share to a relevant company page.')
para('• Reply to every comment in the first 2 hours — LinkedIn\'s algorithm treats comments as a stronger signal than likes.')
para('• If you cross-post: HN title can be sharper — "We made our AI-governance page verify itself with SHA-256 on every load." Lobste.rs: tag as `practices` and `crypto`.')
para('• Do NOT edit the post within the first hour; edits suppress reach.')

# Print char counts so they show up in the run log too
for lbl, txt in [('HOOK_140', HOOK_140), ('HOOK_210', HOOK_210), ('POST_FULL', POST_FULL), ('POST_SHORT', POST_SHORT), ('FIRST_COMMENT', FIRST_COMMENT), ('HASHTAGS', HASHTAGS)]:
    print(f"{lbl}: {len(txt)} chars")

doc.save('exports/a11oy_anatomy_linkedin_post.docx')
print('wrote exports/a11oy_anatomy_linkedin_post.docx')
